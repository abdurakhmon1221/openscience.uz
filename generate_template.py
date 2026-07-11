from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

# Title
title = doc.add_heading('MAQOLA SARLAVHASI (BOSH HARFLAR BILAN, O\'RTADA, QALIN)', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Authors
author = doc.add_paragraph()
author.add_run('Muallifning F.I.Sh.').bold = True
author.add_run('\nIlmiy darajasi, unvoni, ish joyi nomi\nEmail: muallif@mail.com')
author.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('\n')

# Abstract
abstract_heading = doc.add_heading('Annotatsiya (Abstract)', level=1)
abstract_body = doc.add_paragraph(
    'Ushbu qismda maqolaning qisqacha mazmuni, maqsadi, uslubiyati va asosiy natijalari yoziladi. '
    'Annotatsiya hajmi odatda 150-250 so\'zdan iborat bo\'lishi tavsiya etiladi. '
    'Annotatsiya matni aniq, lo\'nda va tushunarli bo\'lishi shart.'
)
abstract_body.paragraph_format.first_line_indent = Inches(0.5)

# Keywords
keywords = doc.add_paragraph()
keywords.add_run('Kalit so\'zlar: ').bold = True
keywords.add_run('kalit so\'z 1, kalit so\'z 2, kalit so\'z 3, kalit so\'z 4, kalit so\'z 5 (5-8 ta so\'z).')

doc.add_paragraph('\n')

# Introduction
intro_heading = doc.add_heading('Kirish (Introduction)', level=1)
doc.add_paragraph(
    'Bu qismda muammoning dolzarbligi, tadqiqotning maqsadi va vazifalari tushuntiriladi. '
    'Nima uchun aynan shu mavzu tanlangani asoslab beriladi.'
).paragraph_format.first_line_indent = Inches(0.5)

# Methodology
doc.add_heading('Tadqiqot Metodologiyasi (Materials and Methods)', level=1)
doc.add_paragraph(
    'Tadqiqot jarayonida qanday usullar va vositalardan foydalanilganligi yoritiladi. '
    'Ilmiy tahlil, tajriba yoki statistik ma\'lumotlar qanday yig\'ilgani tushuntiriladi.'
).paragraph_format.first_line_indent = Inches(0.5)

# Results
doc.add_heading('Natijalar va Muhokama (Results and Discussion)', level=1)
doc.add_paragraph(
    'Tadqiqot natijasida olingan ma\'lumotlar va yutuqlar tahlil qilinadi. '
    'Agar kerak bo\'lsa jadval yoki rasmlar kiritiladi.'
).paragraph_format.first_line_indent = Inches(0.5)

# Conclusion
doc.add_heading('Xulosa (Conclusion)', level=1)
doc.add_paragraph(
    'Maqola so\'ngida tadqiqot bo\'yicha yakuniy xulosa yoziladi. Kelajakda shu yo\'nalishda '
    'amalga oshirilishi mumkin bo\'lgan ishlarga takliflar beriladi.'
).paragraph_format.first_line_indent = Inches(0.5)

# References
doc.add_heading('Foydalanilgan Adabiyotlar (References)', level=1)
ref1 = doc.add_paragraph('1. Muallif F.I. Kitob nomi. Nashriyot, Yil. Betlar.')
ref2 = doc.add_paragraph('2. Muallif F.I. "Maqola nomi". Jurnal nomi, Jild(Son), Yil, Betlar.')

# Ensure public/templates directory exists
os.makedirs('public/templates', exist_ok=True)
doc.save('public/templates/OpenScience_Template.docx')
print("Template created successfully at public/templates/OpenScience_Template.docx")
